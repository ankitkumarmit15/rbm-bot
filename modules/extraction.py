"""
File extraction module: Extract text from PDF, DOCX, DOC, TXT, XLSX, XLS files
"""

import streamlit as st
from pypdf import PdfReader
from docx import Document

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False


def _extract_text_from_excel(uploaded_file) -> str:
    """Read an Excel workbook and convert each row to pipe-separated key:value text."""
    if not OPENPYXL_AVAILABLE:
        st.error("openpyxl not installed — cannot read Excel files.")
        return ""
    try:
        wb = openpyxl.load_workbook(uploaded_file, data_only=True)
        lines = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            headers = [
                str(h).strip() if h is not None else f"Col{i}"
                for i, h in enumerate(rows[0])
            ]
            lines.append(f"=== Sheet: {sheet_name} ===")
            for row in rows[1:]:
                parts = [
                    f"{h}: {str(v).strip()}"
                    for h, v in zip(headers, row)
                    if v is not None and str(v).strip()
                ]
                if parts:
                    lines.append(" | ".join(parts))
        return "\n".join(lines)
    except Exception as e:
        st.error(f"Excel error: {e}")
        return ""


def extract_text_from_file(uploaded_file) -> str:
    """
    Extract text from uploaded file.
    Supports: PDF, DOCX, DOC, TXT, XLSX, XLS
    """
    ext = uploaded_file.name.split(".")[-1].lower()
    text = ""

    if ext == "pdf":
        try:
            for page in PdfReader(uploaded_file).pages:
                text += (page.extract_text() or "") + "\n"
        except Exception as e:
            st.error(f"PDF error: {e}")

    elif ext in ("docx", "doc"):
        try:
            doc = Document(uploaded_file)
            seen_paras: set[str] = set()

            def _para_text(para) -> str:
                t = para.text.strip()
                if t and t not in seen_paras:
                    seen_paras.add(t)
                    return t + "\n"
                return ""

            for para in doc.paragraphs:
                text += _para_text(para)

            for table in doc.tables:
                for row in table.rows:
                    row_parts = []
                    for cell in row.cells:
                        cell_text = " ".join(
                            p.text.strip() for p in cell.paragraphs if p.text.strip()
                        )
                        if cell_text:
                            row_parts.append(cell_text)
                    if row_parts:
                        line = " | ".join(row_parts)
                        if line not in seen_paras:
                            seen_paras.add(line)
                            text += line + "\n"
        except Exception as e:
            st.error(f"DOCX error: {e}")

    elif ext == "txt":
        try:
            raw_bytes = uploaded_file.read()
            text = raw_bytes.decode("utf-8", errors="ignore")
        except Exception as e:
            st.error(f"TXT error: {e}")

    elif ext in ("xlsx", "xls"):
        text = _extract_text_from_excel(uploaded_file)

    return text
